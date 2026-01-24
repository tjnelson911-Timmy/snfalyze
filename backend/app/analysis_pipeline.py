"""
Document Analysis Pipeline

Background processing for document parsing, extraction, and analysis.
"""
import os
import re
import threading
import traceback
from datetime import datetime
from typing import Optional, Dict, Any, List, Tuple
from sqlalchemy.orm import Session

from . import models, schemas, crud
from .database import SessionLocal


# ============================================================================
# DOCUMENT PARSING
# ============================================================================

def parse_pdf_document(file_path: str) -> Dict[str, Any]:
    """
    Parse PDF document and extract text and tables.
    Returns dict with 'pages' (list of page text) and 'tables' (list of extracted tables).
    """
    import pdfplumber

    result = {"pages": [], "tables": [], "metadata": {}}

    try:
        with pdfplumber.open(file_path) as pdf:
            result["metadata"]["page_count"] = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text
                text = page.extract_text() or ""
                result["pages"].append({
                    "page_number": page_num,
                    "text": text,
                    "char_count": len(text)
                })

                # Extract tables
                tables = page.extract_tables()
                for table_idx, table in enumerate(tables):
                    if table and len(table) > 1:
                        # First row as headers
                        headers = [str(h).strip() if h else f"col_{i}" for i, h in enumerate(table[0])]
                        rows = []
                        for row in table[1:]:
                            if any(cell for cell in row):
                                row_dict = {}
                                for i, cell in enumerate(row):
                                    if i < len(headers):
                                        row_dict[headers[i]] = str(cell).strip() if cell else ""
                                rows.append(row_dict)

                        result["tables"].append({
                            "page_number": page_num,
                            "table_index": table_idx,
                            "headers": headers,
                            "rows": rows,
                            "row_count": len(rows)
                        })
    except Exception as e:
        result["error"] = str(e)

    return result


def parse_excel_document(file_path: str) -> Dict[str, Any]:
    """
    Parse Excel document and extract all sheets as tables.
    """
    import openpyxl

    result = {"sheets": [], "tables": [], "metadata": {}}

    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        result["metadata"]["sheet_count"] = len(wb.sheetnames)
        result["metadata"]["sheet_names"] = wb.sheetnames

        for sheet_idx, sheet in enumerate(wb.worksheets):
            sheet_data = {
                "sheet_name": sheet.title,
                "sheet_index": sheet_idx,
                "rows": []
            }

            all_rows = list(sheet.iter_rows(values_only=True))
            if not all_rows:
                continue

            # Detect header row (first row with content)
            headers = None
            data_start = 0
            for i, row in enumerate(all_rows):
                if any(cell for cell in row):
                    headers = [str(cell).strip() if cell else f"col_{j}" for j, cell in enumerate(row)]
                    data_start = i + 1
                    break

            if headers:
                for row in all_rows[data_start:]:
                    if any(cell for cell in row):
                        row_dict = {}
                        for j, cell in enumerate(row):
                            if j < len(headers):
                                row_dict[headers[j]] = cell
                        sheet_data["rows"].append(row_dict)

                sheet_data["headers"] = headers
                sheet_data["row_count"] = len(sheet_data["rows"])
                result["sheets"].append(sheet_data)

                # Also add as a table
                result["tables"].append({
                    "sheet_name": sheet.title,
                    "sheet_index": sheet_idx,
                    "headers": headers,
                    "rows": sheet_data["rows"],
                    "row_count": len(sheet_data["rows"])
                })

        wb.close()
    except Exception as e:
        result["error"] = str(e)

    return result


def parse_word_document(file_path: str) -> Dict[str, Any]:
    """
    Parse Word document and extract text and tables.
    """
    from docx import Document

    result = {"paragraphs": [], "tables": [], "metadata": {}}

    try:
        doc = Document(file_path)

        # Extract paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                result["paragraphs"].append({
                    "index": para_idx,
                    "text": para.text.strip(),
                    "style": para.style.name if para.style else None
                })

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            rows_data = []
            headers = None

            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if row_idx == 0:
                    headers = cells
                else:
                    if any(cells):
                        row_dict = {}
                        for i, cell in enumerate(cells):
                            if i < len(headers):
                                row_dict[headers[i]] = cell
                        rows_data.append(row_dict)

            if headers and rows_data:
                result["tables"].append({
                    "table_index": table_idx,
                    "headers": headers,
                    "rows": rows_data,
                    "row_count": len(rows_data)
                })

        result["metadata"]["paragraph_count"] = len(result["paragraphs"])
        result["metadata"]["table_count"] = len(result["tables"])
    except Exception as e:
        result["error"] = str(e)

    return result


def parse_csv_document(file_path: str) -> Dict[str, Any]:
    """
    Parse CSV document.
    """
    import csv

    result = {"tables": [], "metadata": {}}

    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            # Detect delimiter
            sample = f.read(8192)
            f.seek(0)

            try:
                dialect = csv.Sniffer().sniff(sample)
            except:
                dialect = csv.excel

            reader = csv.reader(f, dialect)
            rows = list(reader)

            if rows:
                headers = [str(h).strip() if h else f"col_{i}" for i, h in enumerate(rows[0])]
                data_rows = []

                for row in rows[1:]:
                    if any(cell for cell in row):
                        row_dict = {}
                        for i, cell in enumerate(row):
                            if i < len(headers):
                                row_dict[headers[i]] = cell.strip() if cell else ""
                        data_rows.append(row_dict)

                result["tables"].append({
                    "headers": headers,
                    "rows": data_rows,
                    "row_count": len(data_rows)
                })

                result["metadata"]["row_count"] = len(data_rows)
                result["metadata"]["column_count"] = len(headers)
    except Exception as e:
        result["error"] = str(e)

    return result


def parse_document(file_path: str, file_type: str) -> Dict[str, Any]:
    """
    Parse a document based on its type.
    """
    parsers = {
        "pdf": parse_pdf_document,
        "xlsx": parse_excel_document,
        "xls": parse_excel_document,
        "docx": parse_word_document,
        "doc": parse_word_document,
        "csv": parse_csv_document,
    }

    parser = parsers.get(file_type.lower())
    if not parser:
        return {"error": f"Unsupported file type: {file_type}"}

    return parser(file_path)


# ============================================================================
# VALUE EXTRACTION PATTERNS
# ============================================================================

# Patterns for extracting financial values
FINANCIAL_PATTERNS = {
    "revenue": [
        r"(?:total\s+)?revenue[:\s]+\$?([\d,]+(?:\.\d+)?)\s*(?:M|million)?",
        r"(?:total\s+)?income[:\s]+\$?([\d,]+(?:\.\d+)?)\s*(?:M|million)?",
    ],
    "ebitda": [
        r"ebitda[:\s]+\$?([\d,]+(?:\.\d+)?)\s*(?:M|million)?",
        r"ebitda\s+of\s+\$?([\d,]+(?:\.\d+)?)",
    ],
    "ebitdar": [
        r"ebitdar[:\s]+\$?([\d,]+(?:\.\d+)?)\s*(?:M|million)?",
        r"ebitdar\s+of\s+\$?([\d,]+(?:\.\d+)?)",
    ],
    "noi": [
        r"(?:net\s+operating\s+income|noi)[:\s]+\$?([\d,]+(?:\.\d+)?)\s*(?:M|million)?",
    ],
    "asking_price": [
        r"(?:asking\s+)?price[:\s]+\$?([\d,]+(?:\.\d+)?)\s*(?:M|million)?",
        r"offered\s+at\s+\$?([\d,]+(?:\.\d+)?)",
        r"list(?:ing)?\s+price[:\s]+\$?([\d,]+(?:\.\d+)?)",
    ],
    "cap_rate": [
        r"cap\s*(?:italization)?\s*rate[:\s]+([\d.]+)\s*%?",
        r"going[- ]in\s+cap[:\s]+([\d.]+)\s*%?",
    ],
    "beds": [
        r"(\d+)\s*(?:licensed\s+)?beds",
        r"bed\s+count[:\s]+(\d+)",
        r"licensed\s+for\s+(\d+)\s+beds",
    ],
    "occupancy": [
        r"occupancy[:\s]+([\d.]+)\s*%?",
        r"([\d.]+)\s*%?\s+occupancy",
        r"average\s+occupancy[:\s]+([\d.]+)",
    ],
}

# Patterns for extracting claims from OMs
CLAIM_PATTERNS = {
    "financial": {
        "revenue_growth": [
            r"revenue\s+(?:has\s+)?(?:grown|increased)\s+(?:by\s+)?([\d.]+)\s*%",
            r"([\d.]+)\s*%\s+revenue\s+growth",
        ],
        "margin": [
            r"(?:operating\s+)?margin\s+of\s+([\d.]+)\s*%",
            r"([\d.]+)\s*%\s+(?:operating\s+)?margin",
        ],
        "profitability": [
            r"net\s+income\s+of\s+\$?([\d,]+(?:\.\d+)?)",
            r"profit(?:ability)?\s+of\s+\$?([\d,]+(?:\.\d+)?)",
        ],
    },
    "operational": {
        "occupancy": [
            r"occupancy\s+(?:rate\s+)?(?:of\s+|at\s+)?([\d.]+)\s*%",
            r"([\d.]+)\s*%\s+occupancy",
            r"stabilized\s+occupancy\s+of\s+([\d.]+)\s*%",
        ],
        "payor_mix": [
            r"([\d.]+)\s*%\s+(?:medicare|medicaid|private)",
            r"(?:medicare|medicaid|private)\s+(?:is\s+|at\s+)?([\d.]+)\s*%",
        ],
        "staffing": [
            r"([\d.]+)\s+hprd",
            r"hours\s+per\s+(?:resident|patient)\s+day[:\s]+([\d.]+)",
        ],
    },
    "quality": {
        "star_rating": [
            r"(\d+)[- ]star\s+(?:overall\s+)?rating",
            r"cms\s+(?:overall\s+)?rating[:\s]+(\d+)",
            r"overall\s+rating\s+of\s+(\d+)\s+star",
        ],
        "survey": [
            r"deficiency[- ]free",
            r"zero\s+deficiencies",
            r"(\d+)\s+(?:total\s+)?deficiencies",
        ],
    },
}


def extract_values_from_text(text: str, patterns: Dict) -> List[Dict]:
    """
    Extract values from text using regex patterns.
    Returns list of extracted values with provenance.
    """
    extracted = []
    text_lower = text.lower()

    for field_type, pattern_list in patterns.items():
        for pattern in pattern_list:
            for match in re.finditer(pattern, text_lower, re.IGNORECASE):
                # Get context around the match
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                snippet = text[start:end]

                extracted.append({
                    "field_type": field_type,
                    "value": match.group(1) if match.groups() else match.group(0),
                    "pattern": pattern,
                    "source_snippet": snippet.strip(),
                    "confidence": 0.7  # Base confidence for regex extraction
                })
                break  # Only take first match per pattern type

    return extracted


def extract_claims_from_text(text: str) -> List[Dict]:
    """
    Extract claims from OM/CIM text for verification.
    """
    claims = []

    for category, type_patterns in CLAIM_PATTERNS.items():
        for claim_type, patterns in type_patterns.items():
            for pattern in patterns:
                for match in re.finditer(pattern, text, re.IGNORECASE):
                    start = max(0, match.start() - 100)
                    end = min(len(text), match.end() + 100)
                    snippet = text[start:end]

                    # Get the full sentence as claim text
                    sentence_start = text.rfind('.', 0, match.start()) + 1
                    sentence_end = text.find('.', match.end())
                    if sentence_end == -1:
                        sentence_end = min(len(text), match.end() + 200)
                    claim_text = text[sentence_start:sentence_end].strip()

                    claims.append({
                        "claim_category": category,
                        "claim_type": claim_type,
                        "claim_text": claim_text[:500],
                        "claimed_value": match.group(1) if match.groups() else None,
                        "source_snippet": snippet.strip(),
                        "confidence": 0.6
                    })

    return claims


# ============================================================================
# ANALYSIS PIPELINE
# ============================================================================

def run_document_analysis(document_id: int, job_id: int):
    """
    Main analysis pipeline for a document.
    Runs in background thread.
    """
    db = SessionLocal()

    try:
        # Get document
        doc = db.query(models.Document).filter(models.Document.id == document_id).first()
        if not doc:
            raise ValueError(f"Document {document_id} not found")

        # Update job status
        job = db.query(models.AnalysisJob).filter(models.AnalysisJob.id == job_id).first()
        if job:
            job.status = "running"
            job.started_at = datetime.utcnow()
            db.commit()

        # Update document parsing status
        doc.parsing_status = "processing"
        db.commit()

        # Get file path
        upload_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "uploads")
        file_path = os.path.join(upload_dir, doc.filename)

        if not os.path.exists(file_path):
            raise ValueError(f"File not found: {file_path}")

        # Parse document
        if job:
            job.progress = 20
            db.commit()

        parsed = parse_document(file_path, doc.file_type)

        if "error" in parsed:
            doc.parsing_status = "error"
            doc.parsing_error = parsed["error"]
            if job:
                job.status = "failed"
                job.error_message = parsed["error"]
            db.commit()
            return

        # Store parsed artifacts
        if job:
            job.progress = 40
            db.commit()

        # Store tables as artifacts
        tables = parsed.get("tables", [])
        for table in tables:
            artifact = models.ParsedArtifact(
                document_id=document_id,
                artifact_type="table",
                page_number=table.get("page_number"),
                content_json=table,
                confidence=0.9
            )
            db.add(artifact)

        # Store text content
        pages = parsed.get("pages", [])
        full_text = ""
        for page in pages:
            artifact = models.ParsedArtifact(
                document_id=document_id,
                artifact_type="text",
                page_number=page.get("page_number"),
                content=page.get("text", ""),
                confidence=1.0
            )
            db.add(artifact)
            full_text += page.get("text", "") + "\n"

        # For Word docs, concatenate paragraphs
        paragraphs = parsed.get("paragraphs", [])
        if paragraphs:
            full_text = "\n".join(p.get("text", "") for p in paragraphs)
            artifact = models.ParsedArtifact(
                document_id=document_id,
                artifact_type="text",
                content=full_text,
                confidence=1.0
            )
            db.add(artifact)

        db.commit()

        if job:
            job.progress = 60
            db.commit()

        # Extract values based on document type
        if doc.doc_type == "om":
            # Extract claims for OM scrubbing
            claims = extract_claims_from_text(full_text)
            for claim_data in claims:
                claim = models.Claim(
                    document_id=document_id,
                    deal_id=doc.deal_id,
                    claim_category=claim_data["claim_category"],
                    claim_type=claim_data["claim_type"],
                    claim_text=claim_data["claim_text"],
                    claimed_value=claim_data.get("claimed_value"),
                    source_snippet=claim_data.get("source_snippet"),
                    confidence=claim_data.get("confidence", 0.5),
                    verification_status="pending"
                )
                db.add(claim)

        # Extract financial values
        extracted = extract_values_from_text(full_text, FINANCIAL_PATTERNS)
        for field_data in extracted:
            field = models.ExtractedField(
                document_id=document_id,
                deal_id=doc.deal_id,
                field_key=field_data["field_type"],
                field_value=field_data["value"],
                field_type="financial",
                source_snippet=field_data.get("source_snippet"),
                confidence=field_data.get("confidence", 0.5),
                extraction_method="regex",
                created_by="system"
            )
            db.add(field)

        db.commit()

        if job:
            job.progress = 80
            db.commit()

        # Update document status
        doc.parsing_status = "completed"
        doc.analyzed = True

        # Generate summary
        summary_parts = []
        if tables:
            summary_parts.append(f"Extracted {len(tables)} table(s)")
        if pages:
            summary_parts.append(f"Processed {len(pages)} page(s)")
        if extracted:
            summary_parts.append(f"Found {len(extracted)} financial value(s)")

        doc.analysis_summary = ". ".join(summary_parts) if summary_parts else "Document processed"

        # Complete job
        if job:
            job.status = "completed"
            job.progress = 100
            job.completed_at = datetime.utcnow()
            job.result = {
                "tables_extracted": len(tables),
                "pages_processed": len(pages) or len(paragraphs),
                "fields_extracted": len(extracted),
                "claims_found": len(claims) if doc.doc_type == "om" else 0
            }

        db.commit()

    except Exception as e:
        traceback.print_exc()
        # Update error status
        try:
            doc = db.query(models.Document).filter(models.Document.id == document_id).first()
            if doc:
                doc.parsing_status = "error"
                doc.parsing_error = str(e)

            job = db.query(models.AnalysisJob).filter(models.AnalysisJob.id == job_id).first()
            if job:
                job.status = "failed"
                job.error_message = str(e)
                job.completed_at = datetime.utcnow()

            db.commit()
        except:
            pass
    finally:
        db.close()


def start_document_analysis(db: Session, document_id: int) -> models.AnalysisJob:
    """
    Start document analysis as a background job.
    Returns the created job.
    """
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc:
        raise ValueError(f"Document {document_id} not found")

    # Create analysis job
    job = models.AnalysisJob(
        deal_id=doc.deal_id,
        document_id=document_id,
        job_type="document_analysis",
        status="pending",
        progress=0
    )
    db.add(job)
    db.commit()
    db.refresh(job)

    # Start background thread
    thread = threading.Thread(
        target=run_document_analysis,
        args=(document_id, job.id),
        daemon=True
    )
    thread.start()

    return job


def start_deal_analysis(db: Session, deal_id: int) -> models.AnalysisJob:
    """
    Start analysis of all documents in a deal.
    Returns the created job.
    """
    deal = db.query(models.Deal).filter(models.Deal.id == deal_id).first()
    if not deal:
        raise ValueError(f"Deal {deal_id} not found")

    # Create analysis job
    job = models.AnalysisJob(
        deal_id=deal_id,
        job_type="deal_analysis",
        status="pending",
        progress=0
    )
    db.add(job)
    db.commit()
    db.refresh(job)

    # Get all unanalyzed documents
    docs = db.query(models.Document).filter(
        models.Document.deal_id == deal_id,
        models.Document.parsing_status.in_(["pending", "error"])
    ).all()

    def run_deal_analysis():
        nonlocal job
        db_inner = SessionLocal()
        try:
            job_inner = db_inner.query(models.AnalysisJob).filter(models.AnalysisJob.id == job.id).first()
            job_inner.status = "running"
            job_inner.started_at = datetime.utcnow()
            db_inner.commit()

            total = len(docs)
            for i, doc in enumerate(docs):
                # Create sub-job for document
                sub_job = models.AnalysisJob(
                    deal_id=deal_id,
                    document_id=doc.id,
                    job_type="document_analysis",
                    status="pending",
                    progress=0
                )
                db_inner.add(sub_job)
                db_inner.commit()
                db_inner.refresh(sub_job)

                # Run analysis synchronously within this thread
                run_document_analysis(doc.id, sub_job.id)

                # Update progress
                job_inner.progress = int((i + 1) / total * 100)
                db_inner.commit()

            job_inner.status = "completed"
            job_inner.progress = 100
            job_inner.completed_at = datetime.utcnow()
            job_inner.result = {"documents_processed": total}
            db_inner.commit()

        except Exception as e:
            traceback.print_exc()
            try:
                job_inner = db_inner.query(models.AnalysisJob).filter(models.AnalysisJob.id == job.id).first()
                job_inner.status = "failed"
                job_inner.error_message = str(e)
                db_inner.commit()
            except:
                pass
        finally:
            db_inner.close()

    # Start background thread
    thread = threading.Thread(target=run_deal_analysis, daemon=True)
    thread.start()

    return job
