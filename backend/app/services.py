"""
Facility Search and Document Analysis Services
"""
import os
import re
import json
import requests
import csv
from io import StringIO
from typing import Optional, List, Dict, Any
from datetime import datetime, timedelta
import threading

# CMS Data - will be cached locally
CMS_CSV_URL = "https://data.cms.gov/provider-data/sites/default/files/resources/66f260a1b66c08187c24fee8d189943b_1767384363/NH_ProviderInfo_Dec2025.csv"
CACHE_FILE = os.path.join(os.path.dirname(__file__), "..", "data", "cms_facilities.json")
CACHE_EXPIRY_HOURS = 24

# In-memory cache
_facilities_cache = None
_cache_lock = threading.Lock()


def get_facilities_data() -> List[Dict]:
    """Get facilities data, loading from cache or downloading if needed"""
    global _facilities_cache

    with _cache_lock:
        # Check if we have in-memory cache
        if _facilities_cache is not None:
            return _facilities_cache

        # Check if we have a valid local cache file
        if os.path.exists(CACHE_FILE):
            try:
                with open(CACHE_FILE, 'r') as f:
                    cache_data = json.load(f)
                    cache_time = datetime.fromisoformat(cache_data.get('timestamp', '2000-01-01'))
                    if datetime.now() - cache_time < timedelta(hours=CACHE_EXPIRY_HOURS):
                        _facilities_cache = cache_data.get('facilities', [])
                        print(f"Loaded {len(_facilities_cache)} facilities from cache")
                        return _facilities_cache
            except Exception as e:
                print(f"Cache read error: {e}")

        # Download fresh data
        print("Downloading CMS facility data...")
        facilities = download_cms_data()

        if facilities:
            _facilities_cache = facilities
            # Save to cache file
            try:
                os.makedirs(os.path.dirname(CACHE_FILE), exist_ok=True)
                with open(CACHE_FILE, 'w') as f:
                    json.dump({
                        'timestamp': datetime.now().isoformat(),
                        'facilities': facilities
                    }, f)
                print(f"Cached {len(facilities)} facilities")
            except Exception as e:
                print(f"Cache write error: {e}")

        return _facilities_cache or []


def download_cms_data() -> List[Dict]:
    """Download and parse CMS nursing home data"""
    facilities = []

    try:
        response = requests.get(CMS_CSV_URL, timeout=60)
        response.raise_for_status()

        # Parse CSV
        csv_content = response.text
        reader = csv.DictReader(StringIO(csv_content))

        for row in reader:
            facility = parse_cms_row(row)
            if facility:
                facilities.append(facility)

        print(f"Downloaded {len(facilities)} facilities from CMS")

    except Exception as e:
        print(f"CMS download error: {e}")
        # Try backup/alternative sources
        facilities = try_backup_sources()

    return facilities


def parse_cms_row(row: Dict) -> Optional[Dict]:
    """Parse a row from CMS CSV into our facility format"""
    try:
        name = row.get('Provider Name', '')
        if not name:
            return None

        # Parse star rating
        star_rating = None
        rating_str = row.get('Overall Rating', '')
        if rating_str and str(rating_str).isdigit():
            star_rating = int(rating_str)

        # Parse beds
        beds = None
        beds_str = row.get('Number of Certified Beds', '')
        if beds_str:
            try:
                beds = int(float(beds_str))
            except:
                pass

        # Parse average residents for occupancy calc
        occupancy = None
        residents_str = row.get('Average Number of Residents per Day', '')
        if residents_str and beds:
            try:
                residents = float(residents_str)
                occupancy = round((residents / beds) * 100, 1) if beds > 0 else None
            except:
                pass

        return {
            "source": "cms",
            "provider_id": row.get('CMS Certification Number (CCN)', ''),
            "name": name.title() if name.isupper() else name,
            "address": (row.get('Provider Address', '') or '').title(),
            "city": (row.get('City/Town', '') or '').title(),
            "state": row.get('State', ''),
            "zip": row.get('ZIP Code', ''),
            "phone": row.get('Telephone Number', ''),
            "licensed_beds": beds,
            "star_rating": star_rating,
            "current_occupancy": occupancy,
            "ownership_type": row.get('Ownership Type', ''),
            "in_hospital": row.get('Provider Resides in Hospital', '') == 'Y',
            "health_rating": row.get('Health Inspection Rating', ''),
            "staffing_rating": row.get('Staffing Rating', ''),
            "quality_rating": row.get('QM Rating', ''),
        }
    except Exception as e:
        return None


def try_backup_sources() -> List[Dict]:
    """Try alternative data sources if main CMS download fails"""
    # Could add more backup sources here
    return []


def search_facilities(query: str, state: Optional[str] = None, limit: int = 20) -> List[Dict[str, Any]]:
    """Search for SNF/nursing home facilities"""
    facilities = get_facilities_data()

    if not facilities:
        return []

    query_lower = query.lower()
    results = []

    for facility in facilities:
        # Filter by state if specified
        if state and facility.get('state', '').upper() != state.upper():
            continue

        # Search by name (fuzzy match)
        name = facility.get('name', '').lower()
        city = facility.get('city', '').lower()

        # Check if query matches name or city
        if query_lower in name or query_lower in city:
            # Calculate relevance score
            score = 0
            if name.startswith(query_lower):
                score = 100
            elif query_lower in name:
                score = 50
            if query_lower in city:
                score += 10

            results.append((score, facility))

    # Sort by relevance score, then by name
    results.sort(key=lambda x: (-x[0], x[1].get('name', '')))

    return [r[1] for r in results[:limit]]


def analyze_document(file_path: str, file_type: str, filename: str) -> Dict[str, Any]:
    """Analyze uploaded documents and extract relevant information"""
    result = {
        "summary": "",
        "document_type": "unknown",
        "extracted_data": {},
        "metrics": {},
        "confidence": 0.0
    }

    text = ""

    try:
        if file_type == "pdf":
            text = extract_pdf_text(file_path)
        elif file_type in ["docx", "doc"]:
            text = extract_docx_text(file_path)
        elif file_type in ["xlsx", "xls", "csv"]:
            text, result["extracted_data"] = extract_spreadsheet_data(file_path, file_type)
        elif file_type in ["png", "jpg", "jpeg"]:
            result["summary"] = "Image file uploaded. Manual review required."
            result["document_type"] = "image"
            return result
    except Exception as e:
        result["summary"] = f"Error extracting text: {str(e)}"
        return result

    if not text:
        result["summary"] = "Could not extract text from document."
        return result

    # Classify document type
    result["document_type"] = classify_document(text, filename)

    # Extract metrics based on document type
    result["metrics"] = extract_metrics(text, result["document_type"])

    # Generate summary
    result["summary"] = generate_summary(text, result["document_type"], result["metrics"])

    result["confidence"] = calculate_confidence(result)

    return result


def extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF"""
    try:
        import PyPDF2
        text = ""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""


def extract_docx_text(file_path: str) -> str:
    """Extract text from Word documents"""
    try:
        from docx import Document
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += "\n" + row_text

        return text
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""


def extract_spreadsheet_data(file_path: str, file_type: str) -> tuple:
    """Extract data from spreadsheets"""
    try:
        import pandas as pd

        if file_type == "csv":
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)

        # Convert to text representation
        text = df.to_string()

        # Extract structured data
        data = {
            "columns": list(df.columns),
            "rows": len(df),
            "preview": df.head(10).to_dict()
        }

        # Try to identify financial columns
        financial_cols = []
        for col in df.columns:
            col_lower = str(col).lower()
            if any(term in col_lower for term in ["revenue", "expense", "income", "ebitda", "noi", "rent", "occupancy", "beds"]):
                financial_cols.append(col)

        if financial_cols:
            data["financial_columns"] = financial_cols
            for col in financial_cols:
                try:
                    if df[col].dtype in ['int64', 'float64']:
                        data[f"{col}_total"] = float(df[col].sum())
                        data[f"{col}_avg"] = float(df[col].mean())
                except:
                    pass

        return text, data
    except Exception as e:
        print(f"Spreadsheet extraction error: {e}")
        return "", {}


def classify_document(text: str, filename: str) -> str:
    """Classify document type based on content and filename"""
    text_lower = text.lower()
    filename_lower = filename.lower()

    # Check filename first
    if any(term in filename_lower for term in ["om", "offering", "memorandum", "teaser"]):
        return "offering_memorandum"
    if any(term in filename_lower for term in ["financial", "p&l", "income", "statement"]):
        return "financials"
    if any(term in filename_lower for term in ["rent", "roll"]):
        return "rent_roll"
    if any(term in filename_lower for term in ["survey", "appraisal"]):
        return "survey"
    if any(term in filename_lower for term in ["license", "certification"]):
        return "license"

    # Check content
    if any(term in text_lower for term in ["offering memorandum", "investment opportunity", "executive summary", "investment highlights"]):
        return "offering_memorandum"
    if any(term in text_lower for term in ["income statement", "profit and loss", "balance sheet", "ebitda", "revenue breakdown"]):
        return "financials"
    if any(term in text_lower for term in ["rent roll", "tenant", "unit mix", "lease"]):
        return "rent_roll"
    if any(term in text_lower for term in ["medicaid", "medicare", "payor mix", "reimbursement"]):
        return "payor_mix"
    if any(term in text_lower for term in ["inspection", "deficiency", "survey", "citation"]):
        return "survey"
    if any(term in text_lower for term in ["star rating", "quality measure", "cms"]):
        return "quality_report"

    return "other"


def extract_metrics(text: str, doc_type: str) -> Dict[str, Any]:
    """Extract financial and operational metrics from text"""
    metrics = {}
    text_lower = text.lower()

    # Currency patterns
    currency_pattern = r'\$[\d,]+(?:\.\d{2})?(?:\s*(?:million|M|thousand|K))?'

    # Extract dollar amounts
    amounts = re.findall(currency_pattern, text)
    if amounts:
        metrics["dollar_amounts_found"] = len(amounts)

    # EBITDAR
    ebitdar_patterns = [
        r'ebitdar[:\s]+\$?([\d,]+(?:\.\d{2})?)',
        r'ebitdar[:\s]+\$?([\d.]+)\s*(?:million|M)',
    ]
    for pattern in ebitdar_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                val = match.group(1).replace(",", "")
                val = float(val)
                if "million" in text_lower[match.start():match.end()+10].lower() or "M" in text[match.start():match.end()+5]:
                    val *= 1000000
                metrics["ebitdar"] = val
                break
            except:
                pass

    # Revenue
    revenue_patterns = [
        r'(?:total\s+)?revenue[:\s]+\$?([\d,]+(?:\.\d{2})?)',
        r'(?:total\s+)?revenue[:\s]+\$?([\d.]+)\s*(?:million|M)',
        r'gross\s+revenue[:\s]+\$?([\d,]+(?:\.\d{2})?)',
    ]
    for pattern in revenue_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                val = match.group(1).replace(",", "")
                val = float(val)
                if "million" in text_lower[match.start():match.end()+10].lower():
                    val *= 1000000
                metrics["revenue"] = val
                break
            except:
                pass

    # Beds
    bed_patterns = [
        r'(\d+)\s*(?:licensed\s+)?beds',
        r'bed\s*(?:count|capacity)[:\s]+(\d+)',
        r'(\d+)\s*bed\s+facility',
    ]
    for pattern in bed_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                metrics["beds"] = int(match.group(1))
                break
            except:
                pass

    # Occupancy
    occ_patterns = [
        r'occupancy[:\s]+(\d+(?:\.\d+)?)\s*%',
        r'(\d+(?:\.\d+)?)\s*%\s*occupancy',
        r'occupancy\s+rate[:\s]+(\d+(?:\.\d+)?)',
    ]
    for pattern in occ_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                metrics["occupancy"] = float(match.group(1))
                break
            except:
                pass

    # Asking price
    price_patterns = [
        r'asking\s*(?:price)?[:\s]+\$?([\d,]+(?:\.\d{2})?)',
        r'list\s*price[:\s]+\$?([\d,]+(?:\.\d{2})?)',
        r'price[:\s]+\$?([\d.]+)\s*(?:million|M)',
    ]
    for pattern in price_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                val = match.group(1).replace(",", "")
                val = float(val)
                if "million" in text_lower[match.start():match.end()+10].lower():
                    val *= 1000000
                metrics["asking_price"] = val
                break
            except:
                pass

    # Cap rate
    cap_patterns = [
        r'cap\s*rate[:\s]+(\d+(?:\.\d+)?)\s*%',
        r'(\d+(?:\.\d+)?)\s*%\s*cap',
    ]
    for pattern in cap_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                metrics["cap_rate"] = float(match.group(1))
                break
            except:
                pass

    # Star rating
    star_patterns = [
        r'(\d)\s*(?:-\s*)?star\s+(?:rating|facility)',
        r'star\s+rating[:\s]+(\d)',
        r'overall\s+rating[:\s]+(\d)',
    ]
    for pattern in star_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                rating = int(match.group(1))
                if 1 <= rating <= 5:
                    metrics["star_rating"] = rating
                    break
            except:
                pass

    return metrics


def generate_summary(text: str, doc_type: str, metrics: Dict) -> str:
    """Generate a human-readable summary of the document"""

    doc_type_names = {
        "offering_memorandum": "Offering Memorandum",
        "financials": "Financial Statement",
        "rent_roll": "Rent Roll",
        "payor_mix": "Payor Mix Report",
        "survey": "Survey/Inspection Report",
        "quality_report": "Quality Report",
        "license": "License/Certification",
        "other": "Document"
    }

    summary_parts = [f"**{doc_type_names.get(doc_type, 'Document')} Analysis**"]

    # Add extracted metrics
    if metrics:
        summary_parts.append("\n**Extracted Metrics:**")

        if "beds" in metrics:
            summary_parts.append(f"- Beds: {metrics['beds']}")
        if "occupancy" in metrics:
            summary_parts.append(f"- Occupancy: {metrics['occupancy']}%")
        if "revenue" in metrics:
            summary_parts.append(f"- Revenue: ${metrics['revenue']:,.0f}")
        if "ebitdar" in metrics:
            summary_parts.append(f"- EBITDAR: ${metrics['ebitdar']:,.0f}")
        if "asking_price" in metrics:
            summary_parts.append(f"- Asking Price: ${metrics['asking_price']:,.0f}")
        if "cap_rate" in metrics:
            summary_parts.append(f"- Cap Rate: {metrics['cap_rate']}%")
        if "star_rating" in metrics:
            summary_parts.append(f"- Star Rating: {metrics['star_rating']}/5")

    # Add content summary based on doc type
    text_lower = text.lower()

    if doc_type == "offering_memorandum":
        sections = []
        if "investment highlights" in text_lower:
            sections.append("Investment Highlights")
        if "property description" in text_lower:
            sections.append("Property Description")
        if "financial" in text_lower:
            sections.append("Financial Overview")
        if "market" in text_lower:
            sections.append("Market Analysis")
        if sections:
            summary_parts.append(f"\n**Sections Found:** {', '.join(sections)}")

    elif doc_type == "financials":
        year_match = re.search(r'(20\d{2})', text)
        if year_match:
            summary_parts.append(f"\n**Period:** {year_match.group(1)}")

    elif doc_type == "survey":
        def_match = re.search(r'(\d+)\s*(?:deficienc|citation)', text_lower)
        if def_match:
            summary_parts.append(f"\n**Deficiencies/Citations:** {def_match.group(1)}")

    # Word count
    word_count = len(text.split())
    summary_parts.append(f"\n**Document Length:** ~{word_count:,} words")

    return "\n".join(summary_parts)


def calculate_confidence(result: Dict) -> float:
    """Calculate confidence score for the analysis"""
    score = 0.0

    if result["document_type"] != "unknown":
        score += 0.3

    metrics_count = len(result.get("metrics", {}))
    score += min(0.4, metrics_count * 0.1)

    if result.get("summary"):
        score += 0.2

    if result.get("extracted_data"):
        score += 0.1

    return min(1.0, score)


# Pre-load facilities data on module import (in background)
def _preload_data():
    try:
        get_facilities_data()
    except Exception as e:
        print(f"Preload error: {e}")

# Start preload in background thread
_preload_thread = threading.Thread(target=_preload_data, daemon=True)
_preload_thread.start()
